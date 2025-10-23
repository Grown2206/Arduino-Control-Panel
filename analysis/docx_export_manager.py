#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
analysis/docx_export_manager.py
Exportiert komplette Berichte als DOCX mit allen Diagrammen und Formatierungen
"""

import base64
import io
from datetime import datetime
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxExportManager:
    """Manager für umfassende DOCX-Exporte mit Diagrammen."""
    
    @staticmethod
    def export_single_report(run_details: Dict[str, Any], analysis: Dict[str, Any], 
                            charts_base64: List[str], file_path: str) -> bool:
        """
        Exportiert einen Einzelbericht als DOCX mit allen Diagrammen.
        
        Args:
            run_details: Dictionary mit Testlauf-Details
            analysis: Dictionary mit Analyse-Ergebnissen
            charts_base64: Liste von Base64-encodierten Diagrammen
            file_path: Ziel-Dateipfad
            
        Returns:
            True bei Erfolg, False bei Fehler
        """
        if not DOCX_AVAILABLE:
            print("❌ python-docx nicht installiert!")
            return False
        
        try:
            doc = Document()
            
            # Titel
            title = doc.add_heading(f"Testbericht: {run_details.get('name', 'Unbekannt')}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Untertitel
            subtitle = doc.add_paragraph(f"Testlauf ID: {run_details.get('id')}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_format = subtitle.runs[0]
            subtitle_format.font.size = Pt(12)
            subtitle_format.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
            
            # === ZUSAMMENFASSUNG ===
            doc.add_heading('Zusammenfassung', 1)
            
            summary_data = [
                ('Sequenz', run_details.get('sequence_name', '-')),
                ('Startzeit', run_details.get('start_time', '-')),
                ('Endzeit', run_details.get('end_time', '-')),
                ('Dauer', f"{run_details.get('duration', 0):.2f} Sekunden"),
                ('Zyklen', str(run_details.get('cycles', 0))),
                ('Status', run_details.get('status', '-')),
            ]
            
            table = doc.add_table(rows=len(summary_data), cols=2)
            table.style = 'Light Grid Accent 1'
            
            for i, (key, value) in enumerate(summary_data):
                row = table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = str(value)
                # Fette Schrift für Keys
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # === ANALYSE-ERGEBNISSE ===
            if analysis:
                doc.add_heading('Analyse-Ergebnisse', 1)
                
                quality = analysis.get('quality_metrics', {})
                perf_rating = analysis.get('performance_rating', {})
                cycle_analysis = analysis.get('cycle_analysis', {})

                # Performance Rating
                doc.add_heading('Performance Rating', 2)
                perf_table = doc.add_table(rows=3, cols=2)
                perf_table.style = 'Light Grid Accent 1'
                
                perf_data = [
                    ('Gesamt-Score', f"{quality.get('overall_score', 0):.1f}"),
                    ('Rating', perf_rating.get('rating', 'N/A')),
                    ('Ergebnis', perf_rating.get('pass_fail', 'N/A')),
                ]
                for i, (key, value) in enumerate(perf_data):
                    row = perf_table.rows[i]
                    row.cells[0].text = key
                    row.cells[1].text = value
                    row.cells[0].paragraphs[0].runs[0].font.bold = True

                doc.add_paragraph()

                # Qualitätsmetriken
                doc.add_heading('Qualitätsmetriken', 2)
                qual_table = doc.add_table(rows=4, cols=2)
                qual_table.style = 'Light Grid Accent 1'
                
                qual_data = [
                    ('Konsistenz', f"{quality.get('consistency_score', 0):.1f} %"),
                    ('Zuverlässigkeit', f"{quality.get('reliability_score', 0):.1f} %"),
                    ('Präzision (Jitter)', f"{quality.get('jitter_score', 0):.1f} %"),
                    ('Anomalierate', f"{quality.get('anomaly_rate', 0):.1f} %"),
                ]
                for i, (key, value) in enumerate(qual_data):
                    row = qual_table.rows[i]
                    row.cells[0].text = key
                    row.cells[1].text = value
                    row.cells[0].paragraphs[0].runs[0].font.bold = True

                doc.add_paragraph()

                if cycle_analysis:
                    doc.add_heading('Zykluszeiten', 2)
                    
                    cycle_table = doc.add_table(rows=5, cols=2)
                    cycle_table.style = 'Light Grid Accent 1'
                    
                    cycle_data = [
                        ('Durchschnitt', f"{cycle_analysis.get('avg', 0):.2f} ms"),
                        ('Minimum', f"{cycle_analysis.get('min', 0):.2f} ms"),
                        ('Maximum', f"{cycle_analysis.get('max', 0):.2f} ms"),
                        ('Standardabweichung', f"{cycle_analysis.get('std', 0):.2f} ms"),
                        ('Stabilität', f"{cycle_analysis.get('stability', 0):.1f} %"),
                    ]
                    
                    for i, (key, value) in enumerate(cycle_data):
                        row = cycle_table.rows[i]
                        row.cells[0].text = key
                        row.cells[1].text = value
                        row.cells[0].paragraphs[0].runs[0].font.bold = True
                    
                    # Anomalien
                    anomalies = cycle_analysis.get('anomalies', [])
                    if anomalies:
                        doc.add_paragraph()
                        anomaly_p = doc.add_paragraph(f'⚠️ Anomalien gefunden: {len(anomalies)}')
                        anomaly_p.runs[0].font.color.rgb = RGBColor(231, 76, 60)  # Rot
                        anomaly_p.runs[0].font.bold = True
                    else:
                        doc.add_paragraph()
                        ok_p = doc.add_paragraph('✓ Keine Anomalien gefunden')
                        ok_p.runs[0].font.color.rgb = RGBColor(46, 204, 113)  # Grün
                
                doc.add_paragraph()
            
            # === SENSOR-DATEN ===
            sensors = run_details.get('log', {}).get('sensors', {})
            if sensors:
                doc.add_heading('Sensor-Daten', 1)
                
                sensor_table = doc.add_table(rows=1, cols=4)
                sensor_table.style = 'Light Grid Accent 1'
                
                # Header
                header_cells = sensor_table.rows[0].cells
                header_cells[0].text = 'Sensor'
                header_cells[1].text = 'Min'
                header_cells[2].text = 'Max'
                header_cells[3].text = 'Durchschnitt'
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                
                # Temperatur
                if 'temp' in sensors:
                    temp = sensors['temp']
                    row = sensor_table.add_row().cells
                    row[0].text = 'Temperatur (°C)'
                    row[1].text = f"{temp.get('min', 0):.1f}"
                    row[2].text = f"{temp.get('max', 0):.1f}"
                    row[3].text = f"{temp.get('avg', 0):.1f}"
                
                # Luftfeuchtigkeit
                if 'humid' in sensors:
                    humid = sensors['humid']
                    row = sensor_table.add_row().cells
                    row[0].text = 'Luftfeuchtigkeit (%)'
                    row[1].text = f"{humid.get('min', 0):.1f}"
                    row[2].text = f"{humid.get('max', 0):.1f}"
                    row[3].text = f"{humid.get('avg', 0):.1f}"
                
                doc.add_paragraph()
            
            # === DIAGRAMME ===
            if charts_base64:
                doc.add_page_break()
                doc.add_heading('Diagramme', 1)
                
                for i, chart_b64 in enumerate(charts_base64, 1):
                    try:
                        # Base64 zu Bild
                        image_data = base64.b64decode(chart_b64)
                        image_stream = io.BytesIO(image_data)
                        
                        # Bild einfügen (6 Zoll Breite)
                        doc.add_picture(image_stream, width=Inches(6))
                        
                        # Zentrieren
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        doc.add_paragraph()
                        
                    except Exception as e:
                        print(f"⚠️ Fehler beim Einfügen von Diagramm {i}: {e}")
            
            # === EVENT-LOG (Optional, nur erste 100) ===
            events = run_details.get('log', {}).get('events', [])
            if events:
                doc.add_page_break()
                doc.add_heading('Event-Log (erste 100 Einträge)', 1)
                
                event_table = doc.add_table(rows=1, cols=4)
                event_table.style = 'Light Grid Accent 1'
                
                # Header
                header_cells = event_table.rows[0].cells
                header_cells[0].text = 'Zeit (ms)'
                header_cells[1].text = 'Pin'
                header_cells[2].text = 'Aktion'
                header_cells[3].text = 'Zyklus'
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                
                # Events (max 100)
                for event in events[:100]:
                    row = event_table.add_row().cells
                    row[0].text = f"{event.get('time', 0):.2f}"
                    row[1].text = str(event.get('pin', '-'))
                    row[2].text = str(event.get('action', '-'))
                    row[3].text = str(event.get('cycle', '-'))
            
            # === FOOTER ===
            doc.add_paragraph()
            footer = doc.add_paragraph(
                f"Erstellt am: {datetime.now().strftime('%d.%m.%Y um %H:%M:%S')}"
            )
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Speichern
            doc.save(file_path)
            return True
            
        except Exception as e:
            print(f"❌ DOCX-Export fehlgeschlagen: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    @staticmethod
    def export_comparison_report(run_details_list: List[Dict[str, Any]], 
                                analysis_results: List[Dict[str, Any]],
                                charts_base64: List[str], file_path: str) -> bool:
        """
        Exportiert einen Vergleichsbericht als DOCX.
        
        Args:
            run_details_list: Liste von Testlauf-Details
            analysis_results: Liste von Analyse-Ergebnissen
            charts_base64: Liste von Base64-encodierten Diagrammen
            file_path: Ziel-Dateipfad
            
        Returns:
            True bei Erfolg, False bei Fehler
        """
        if not DOCX_AVAILABLE:
            print("❌ python-docx nicht installiert!")
            return False
        
        try:
            doc = Document()
            
            # Titel
            title = doc.add_heading(f"Vergleichsbericht: {len(run_details_list)} Testläufe", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # === ÜBERSICHT ===
            doc.add_heading('Übersicht', 1)
            
            # Tabelle mit allen Tests
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            headers = ['ID', 'Name', 'Zyklen', 'Ø Zeit (ms)', 'Stabilität (%)', 
                       'Anomalien', 'Gesamt-Score', 'Rating']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Daten
            for details, analysis in zip(run_details_list, analysis_results):
                cycle_analysis = analysis.get('cycle_analysis', {})
                quality = analysis.get('quality_metrics', {})
                perf_rating = analysis.get('performance_rating', {})
                
                row = table.add_row().cells
                row[0].text = str(details.get('id', '-'))
                row[1].text = details.get('name', '-')
                row[2].text = str(details.get('cycles', 0))
                row[3].text = f"{cycle_analysis.get('avg', 0):.2f}"
                row[4].text = f"{cycle_analysis.get('stability', 0):.1f}"
                row[5].text = str(len(cycle_analysis.get('anomalies', [])))
                row[6].text = f"{quality.get('overall_score', 0):.1f}"
                row[7].text = perf_rating.get('rating', 'N/A')
            
            doc.add_paragraph()
            
            # === VERGLEICHS-DIAGRAMME ===
            if charts_base64:
                doc.add_page_break()
                doc.add_heading('Grafischer Vergleich', 1)
                
                for chart_b64 in charts_base64:
                    try:
                        image_data = base64.b64decode(chart_b64)
                        image_stream = io.BytesIO(image_data)
                        doc.add_picture(image_stream, width=Inches(6.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"⚠️ Fehler beim Einfügen von Vergleichs-Diagramm: {e}")
            
            # === EINZELNE TESTLÄUFE ===
            doc.add_page_break()
            doc.add_heading('Detaillierte Testlauf-Daten', 1)
            
            for i, (details, analysis) in enumerate(zip(run_details_list, analysis_results), 1):
                doc.add_heading(f"Test {i}: {details.get('name')}", 2)
                
                # Details-Tabelle
                detail_table = doc.add_table(rows=4, cols=2)
                detail_table.style = 'Light Grid Accent 1'
                
                detail_data = [
                    ('Start', details.get('start_time', '-')),
                    ('Dauer', f"{details.get('duration', 0):.2f} s"),
                    ('Zyklen', str(details.get('cycles', 0))),
                    ('Status', details.get('status', '-')),
                ]
                
                for j, (key, value) in enumerate(detail_data):
                    row = detail_table.rows[j]
                    row.cells[0].text = key
                    row.cells[1].text = value
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                
                doc.add_paragraph()
            
            # Footer
            footer = doc.add_paragraph(
                f"Vergleichsbericht erstellt am: {datetime.now().strftime('%d.%m.%Y um %H:%M:%S')}"
            )
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Speichern
            doc.save(file_path)
            return True
            
        except Exception as e:
            print(f"❌ Vergleichsbericht DOCX-Export fehlgeschlagen: {e}")
            import traceback
            traceback.print_exc()
            return False